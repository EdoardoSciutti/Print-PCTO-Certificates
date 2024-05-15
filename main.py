import os
import sys
from dotenv import load_dotenv
import tkinter as tk
from tkinter import ttk
from docx import Document
import shutil
import mysql.connector
from datetime import date

conn = None

def informazioniScript():
    os.system('cls' if os.name == 'nt' else 'clear')
    if getattr(sys, 'frozen', False):
        percorso = sys.executable
    else:
        percorso = __file__
    informazioniFile = os.stat(percorso)
    dimensioneFile = (informazioniFile.st_size/1024)/1024 # sono pigro
    print("Autore: Sciutti Edoardo")
    print(f"Dimensione del file: {dimensioneFile:.2f} MB\n")
    print("Registro console, tieni d'occhio le istruzioni")
    procedi = input("Premi invio per procedere...\n")
    connessioneDatabase()
# La funzione informazioniScript pulisce la console, calcola la dimensione del file dello script, 
# stampa le informazioni dello script e attende l'input dell'utente per procedere.

def configurazione():
    load_dotenv()
    hostDatabase = os.getenv('DB_HOST', 'ERRORE')
    nomeDatabase = os.getenv('DB_NAME', 'ERRORE')
    utenteDatabase = os.getenv('DB_USER', 'ERRORE')
    passwordDatabase = os.getenv('DB_PASS', 'ERRORE')
    if hostDatabase == 'ERRORE' or nomeDatabase == 'ERRORE' or utenteDatabase == 'ERRORE' or passwordDatabase == 'ERRORE':
        print("Errore:\n .env non configurato correttamente")
        print("Per avere info su come configurare il file .env leggi il README.md")
        esci = input("Premi invio per chiudere...\n")
        exit()
    return hostDatabase, utenteDatabase, passwordDatabase, nomeDatabase
# La funzione configurazione carica le variabili d'ambiente dal file .env, controlla se sono state impostate correttamente,
# stampa un messaggio di errore se non lo sono, e restituisce le variabili d'ambiente.

def connessioneDatabase():
    global connessione
    hostDatabase, utenteDatabase, passwordDatabase, nomeDatabase = configurazione()
    try:
        connessione = mysql.connector.connect(
            host = hostDatabase,
            user = utenteDatabase,
            password = passwordDatabase,
            database = nomeDatabase
        )
        print(f"Connessione al database {nomeDatabase} avvenuta con successo\n")
        applicazione()
    except mysql.connector.Error as errore:
        print(f"Errore durante la connessione al database:\n --> {errore}\n")
        esci = input("Premi invio per chiudere...\n")
        exit()
# La funzione connessioneDatabase stabilisce una connessione con il database utilizzando le informazioni ottenute dalla funzione di configurazione.
# Se la connessione è riuscita, chiama la funzione applicazione. In caso di errore, stampa l'errore e termina l'esecuzione.

def gestioneQuery(*opzioniQuery):
    global connessione
    cursore = connessione.cursor()

    if opzioniQuery[0] == 1:
        query = """ SELECT description
                    FROM `ca_school_years`"""
        cursore.execute(query)
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]) for tupla in risultatoQuery]
        return sorted(risultatoQuery)

    elif opzioniQuery[0] == 2:
        if opzioniQuery[1] == "Selezione AS":
            return []
        query = """ SELECT name
                    FROM `ca_school_classes`
                    JOIN ca_school_years ON ca_school_classes.school_year_id = ca_school_years.id
                    WHERE ca_school_years.description = %s;"""
        cursore.execute(query, (opzioniQuery[1],))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]) for tupla in risultatoQuery]
        if len(risultatoQuery) == 0:
            return ["Nessuna Classe Trovata"]
        return sorted(risultatoQuery)

    elif opzioniQuery[0] == 3:
        if opzioniQuery[1] == "Selezione Classe" or opzioniQuery[2] == "Selezione AS":
            return []
        query = """ SELECT CONCAT(ca_users.surname, " ", ca_users.name) 
                    FROM `ca_users`
                    JOIN ca_role_user on ca_users.id = ca_role_user.user_id
                    JOIN ca_frequented_classes on ca_frequented_classes.user_id = ca_users.id
                    JOIN ca_school_classes on ca_school_classes.id = ca_frequented_classes.school_class_id
                    JOIN ca_school_years on ca_school_classes.school_year_id = ca_school_years.id
                    WHERE ca_role_user.role_id = 1 AND ca_school_classes.name = %s AND ca_school_years.description = %s;"""
        cursore.execute(query, (opzioniQuery[2], opzioniQuery[1]))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        if len(risultatoQuery) == 0:
            return ["Nessun Alunno Trovato"]
        return sorted(risultatoQuery)

    elif opzioniQuery[0] == 4:
        query = """ SELECT ca_courses.title FROM ca_courses
                    JOIN ca_registrations ON ca_registrations.course_id = ca_courses.id
                    JOIN ca_users ON ca_users.id = ca_registrations.user_id
                    WHERE ca_users.name = %s AND ca_users.surname = %s;"""
        cursore.execute(query, (opzioniQuery[2], opzioniQuery[1]))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sorted(risultatoQuery)

    elif opzioniQuery[0] == 5:
        query = """ SELECT HOUR(TIMEDIFF(ca_presences.exited_at, ca_presences.joined_at)) FROM ca_presences
                    JOIN ca_users ON ca_users.id = ca_presences.user_id
                    JOIN ca_activities ON ca_activities.id = ca_presences.activity_id
                    JOIN ca_courses ON ca_courses.id = ca_activities.course_id
                    WHERE ca_users.name = %s AND ca_users.surname = %s AND ca_courses.title = %s"""
        cursore.execute(query, (opzioniQuery[2], opzioniQuery[1], opzioniQuery[3]))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sum(int(item) for item in risultatoQuery if item != 'None')

    elif opzioniQuery[0] == 6:
        query = """ SELECT HOUR(TIMEDIFF(ca_activities.ended_at, ca_activities.started_at)) FROM ca_activities
                    JOIN ca_courses ON ca_courses.id = ca_activities.course_id
                    WHERE ca_courses.title = %s"""
        cursore.execute(query, (opzioniQuery[1],))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sum(int(item) for item in risultatoQuery)

    elif opzioniQuery[0] == 7:
        query = """ SELECT CONCAT(ca_users.surname, " ", ca_users.name) FROM ca_users
                    JOIN ca_courses ON ca_courses.user_id = ca_users.id
                    WHERE ca_courses.title = %s"""
        cursore.execute(query, (opzioniQuery[1],))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sorted(risultatoQuery)

    elif opzioniQuery[0] == 8:
        query = """ SELECT ca_archive_courses.title FROM ca_archive_courses
                    JOIN ca_archive_registrations ON ca_archive_registrations.course_id = ca_archive_courses.id
                    JOIN ca_users ON ca_users.id = ca_archive_registrations.user_id
                    WHERE ca_users.name = %s AND ca_users.surname = %s;"""
        cursore.execute(query, (opzioniQuery[2], opzioniQuery[1]))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sorted(risultatoQuery)

    elif opzioniQuery[0] == 9:
        query = """ SELECT HOUR(TIMEDIFF(ca_archive_presences.exited_at, ca_archive_presences.joined_at)) FROM ca_archive_presences
                    JOIN ca_users ON ca_users.id = ca_archive_presences.user_id
                    JOIN ca_archive_activities ON ca_archive_activities.id = ca_archive_presences.activity_id
                    JOIN ca_archive_courses ON ca_archive_courses.id = ca_archive_activities.course_id
                    WHERE ca_users.name = %s AND ca_users.surname = %s AND ca_archive_courses.title = %s"""
        cursore.execute(query, (opzioniQuery[2], opzioniQuery[1], opzioniQuery[3]))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sum(int(item) for item in risultatoQuery if item != 'None')

    elif opzioniQuery[0] == 10:
        query = """ SELECT HOUR(TIMEDIFF(ca_archive_activities.ended_at, ca_archive_activities.started_at)) FROM ca_archive_activities
                    JOIN ca_archive_courses ON ca_archive_courses.id = ca_archive_activities.course_id
                    WHERE ca_archive_courses.title = %s"""
        cursore.execute(query, (opzioniQuery[1],))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sum(int(item) for item in risultatoQuery)

    elif opzioniQuery[0] == 11:
        query = """ SELECT CONCAT(ca_users.surname, " ", ca_users.name) FROM ca_users
                    JOIN ca_archive_courses ON ca_archive_courses.user_id = ca_users.id
                    WHERE ca_archive_courses.title = %s"""
        cursore.execute(query, (opzioniQuery[1],))
        risultatoQuery = cursore.fetchall()
        risultatoQuery = [str(tupla[0]).title().lstrip() for tupla in risultatoQuery]
        return sorted(risultatoQuery)
# La funzione gestioneQuery esegue diverse query SQL sul database in base alle opzioni fornite.
# Le opzioni determinano quale query viene eseguita e quali parametri vengono utilizzati.
# I risultati delle query vengono restituiti come liste di stringhe.

def stampaAttestati(*opzioniQuery):
    def riempiTabella(documento, supporto, studente):
        oreTotali = 0
        tabella = documento.tables[0]
        for elemento in gestioneQuery(4 + supporto, studente.split()[0], studente.split()[1]):
            riga = tabella.add_row()
            for indice, cella in enumerate(riga.cells):
                if indice == 0:
                    cella.text = elemento
                if indice == 1:
                    cella.text = str(gestioneQuery(5 + supporto, studente.split()[0], studente.split()[1], elemento))
                    oreTotali += int(cella.text)
                if indice == 2:
                    cella.text = str(gestioneQuery(6 + supporto, elemento))
                if indice == 3:
                    cella.text = gestioneQuery(7 + supporto, elemento)
        return oreTotali

    def sostituisciSegnaposti(documento, studente, classe, anno, oreTotali):
        for paragrafo in documento.paragraphs:
            for esecuzione in paragrafo.runs:
                if 'Placeholdernome' in esecuzione.text:
                    esecuzione.text = esecuzione.text.replace("Placeholdernome", studente)
                if 'Placeholderclasse' in esecuzione.text:
                    esecuzione.text = esecuzione.text.replace("Placeholderclasse", classe)
                if 'Placeholderanno' in esecuzione.text:
                    esecuzione.text = esecuzione.text.replace("Placeholderanno", anno)
                if 'Placeholderdata' in esecuzione.text:
                    esecuzione.text = esecuzione.text.replace("Placeholderdata", date.today().strftime("%d-%m-%Y"))
                if 'Placeholderore' in esecuzione.text:
                    esecuzione.text = esecuzione.text.replace("Placeholderore", str(oreTotali))

    if opzioniQuery[0] == 0:
        shutil.copyfile('base.docx', f'Risultati/{opzioniQuery[1]}.docx')
        documento = Document(f'Risultati/{opzioniQuery[1]}.docx')
        supporto = 4 if opzioniQuery[4] else 0
        oreTotali = riempiTabella(documento, supporto, opzioniQuery[1])
        sostituisciSegnaposti(documento, opzioniQuery[1], opzioniQuery[2], opzioniQuery[3], oreTotali)
        documento.save(f'Risultati/{opzioniQuery[1]}.docx')

    if opzioniQuery[0] == 1:
        os.makedirs(f"Risultati/{opzioniQuery[2]}" + " - " + f"{opzioniQuery[3]}", exist_ok=True)
        directory = f"Risultati/{opzioniQuery[2]}" + " - " + f"{opzioniQuery[3]}"
        for studente in opzioniQuery[1]:
            shutil.copyfile('base.docx', directory + f'/{studente}.docx')
            documento = Document(directory + f'/{studente}.docx')
            supporto = 4 if opzioniQuery[4] else 0
            oreTotali = riempiTabella(documento, supporto, studente)
            sostituisciSegnaposti(documento, studente, opzioniQuery[2], opzioniQuery[3], oreTotali)
            documento.save(directory + f'/{studente}.docx')
# La funzione stampaAttestati genera attestati in formato .docx per uno o più studenti.
# Gli attestati sono basati su un modello di documento (base.docx) e contengono informazioni specifiche dello studente,
# come il nome, la classe, l'anno e il totale delle ore.
# Le informazioni specifiche dello studente sono ottenute eseguendo diverse query sul database.
# Gli attestati generati vengono salvati in una directory Risultati.

def applicazione():
    def allaChiusura():
        global conn
        if conn is not None:
            conn.close()
        print("Connessione al database interrotta correttamente")
        finestraPrincipale.destroy()
        esc = input("Premi invio per chiudere...\n")

    def interruttoreSchermoIntero(event=None):
        finestraPrincipale.attributes('-fullscreen', not finestraPrincipale.attributes('-fullscreen'))

    def alCambioVar1(*args):
        global nomeAnno
        nomeAnno = varAnno.get()
        menuACascata2['values'] = gestioneQuery(2, nomeAnno)
        menuACascata2.set('Selezione Classe')

    def alCambioVar2(*args):
        global nomeClasse
        nomeAnno = varAnno.get()
        nomeClasse = varClasse.get()
        menuACascata3['values'] = gestioneQuery(3, nomeAnno, nomeClasse)
        menuACascata3.set('Selezione Alunno')

    def alCambioVar4(*args):
        if varAnnoAttuale.get() != "Selezione AS Attuale":
            tuttiValori = gestioneQuery(1)
            valoriFiltrati = [valore for valore in tuttiValori if valore <= varAnnoAttuale.get()]
            menuACascata1['values'] = valoriFiltrati
            menuACascata1.config(state="readonly")
            menuACascata1.set("Selezione AS")
        else:
            menuACascata1.config(state="disabled")

    def aggiornaMenuACascata2(*args):
        if varAnno.get() != "Selezione AS":
            menuACascata2.config(state="readonly")
        else:
            menuACascata2.config(state="disabled")

    def aggiornaMenuACascata3(*args):
        if varClasse.get() != "Selezione Classe" and varAnno.get() != "Selezione AS":
            menuACascata3.config(state="readonly")
        else:
            menuACascata3.config(state="disabled")

    def alClickBottone2():
        ASattuale = False if menuACascata1.get() == menuACascata4.get() else True
        stampaAttestati(1, gestioneQuery(3, nomeAnno, nomeClasse), varClasse.get(), varAnno.get(), ASattuale)

    def alClickBottone3():
        ASattuale = False if menuACascata1.get() == menuACascata4.get() else True
        stampaAttestati(0, varAlunno.get(), varClasse.get(), varAnno.get(), ASattuale)

    def controllaMenuACascata(*args):
        if varAnno.get() != "Selezione AS" and varClasse.get() != "Selezione Classe" and varClasse.get() != "Nessuna Classe Trovata":
            bottone2.config(state="normal")
        else:
            bottone2.config(state="disabled")

        if varAnno.get() != "Selezione AS" and varClasse.get() != "Selezione Classe" and varClasse.get() != "Nessuna Classe Trovata" and varAlunno.get() != "Selezione Alunno" and varAlunno.get() != "Nessun Alunno Trovato":
            bottone3.config(state="normal")
        else:
            bottone3.config(state="disabled")

    finestraPrincipale = tk.Tk()
    finestraPrincipale.title("Certificazioni PCTO")
    larghezzaSchermo = finestraPrincipale.winfo_screenwidth()
    altezzaSchermo = finestraPrincipale.winfo_screenheight()
    finestraPrincipale.iconbitmap('icona.ico')
    larghezza = int(larghezzaSchermo // 1.5)
    altezza = altezzaSchermo // 2
    finestraPrincipale.geometry(f'{larghezza}x{altezza}')
    finestraPrincipale.resizable(False, False)
    finestraPrincipale.bind("<F11>", interruttoreSchermoIntero)
    finestraPrincipale.protocol("WM_DELETE_WINDOW", allaChiusura)

    etichetta1 = tk.Label(finestraPrincipale, text="Anno Scolastico", font=("Helvetica", 16))
    etichetta1.grid(row=0, column=1)
    etichetta2 = tk.Label(finestraPrincipale, text="Classe", font=("Helvetica", 16))
    etichetta2.grid(row=0, column=3)
    etichetta3 = tk.Label(finestraPrincipale, text="Alunni", font=("Helvetica", 16))
    etichetta3.grid(row=0, column=5)
    etichetta4 = tk.Label(finestraPrincipale, text="Seleziona correttamente l'Anno\nScolatico attuale altrimenti\nle Query non funzioneranno", font=("Helvetica", 14), fg="red")
    etichetta4.grid(row=2, column=1)

    varAnno = tk.StringVar(finestraPrincipale)
    varClasse = tk.StringVar(finestraPrincipale)
    varAlunno = tk.StringVar(finestraPrincipale)
    varAnnoAttuale = tk.StringVar(finestraPrincipale)

    menuACascata1 = ttk.Combobox(finestraPrincipale, textvariable=varAnno, values=gestioneQuery(1), font=("Helvetica", 16), state="disabled", width=23)
    menuACascata1.set("Selezione AS")
    menuACascata1.grid(row=1, column=1)
    varAnno.trace_add("write", alCambioVar1)
    varAnno.trace_add("write", aggiornaMenuACascata2)
    varAnno.trace_add("write", controllaMenuACascata)

    menuACascata2 = ttk.Combobox(finestraPrincipale, textvariable=varClasse, values=[], font=("Helvetica", 16), state="disabled", width=23)
    menuACascata2.set("Selezione Classe")
    menuACascata2.grid(row=1, column=3)
    varClasse.trace_add("write", alCambioVar2)
    varClasse.trace_add("write", aggiornaMenuACascata3)
    varClasse.trace_add("write", controllaMenuACascata)

    menuACascata3 = ttk.Combobox(finestraPrincipale, textvariable=varAlunno, values=[], font=("Helvetica", 16), state="disabled", width=23)
    menuACascata3.set("Selezione Alunno")
    menuACascata3.grid(row=1, column=5)
    varAlunno.trace_add("write", controllaMenuACascata)

    menuACascata4 = ttk.Combobox(finestraPrincipale, textvariable=varAnnoAttuale, values=gestioneQuery(1), font=("Helvetica", 16), state = "readonly", width = 23)
    menuACascata4.set("Selezione AS Attuale")
    menuACascata4.grid(row=3, column=1)
    varAnnoAttuale.trace_add("write", alCambioVar4)

    bottone2 = tk.Button(finestraPrincipale, text="Stampa Attestati\ndella Classe", font=("Helvetica", 16), state="disabled", width=23, command=alClickBottone2)
    bottone2.grid(row=2, column=3)

    bottone3 = tk.Button(finestraPrincipale, text="Stampa Attestato\ndell'Alunno", font=("Helvetica", 16), command=alClickBottone3, state="disabled", width=23)
    bottone3.grid(row=2, column=5)

    for i in range(7):
        finestraPrincipale.grid_columnconfigure(i, weight=1)
    for i in range(4):
        finestraPrincipale.grid_rowconfigure(i, weight=1)

    finestraPrincipale.mainloop()
# La funzione applicazione crea una interfaccia utente grafica (GUI) per l'applicazione. 
# Questa GUI include vari elementi come etichette, menu a cascata e pulsanti, 
# e gestisce diversi eventi come il cambio di selezione nei menu a cascata e il click sui pulsanti. 
# Alla fine, entra in un ciclo principale per gestire gli eventi dell'interfaccia utente.

def main():
    informazioniScript()
# La funzione main() chiama le funzioni per visualizzare le informazioni dello script
  
if __name__ == "__main__":
    main()